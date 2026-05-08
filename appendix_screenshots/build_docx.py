"""
Build Appendix DOCX v2 — Times New Roman, 12pt, page numbers from 22.
More screenshots, more visualizations, code stays concise.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SS_DIR = os.path.join(PROJECT, "appendix_screenshots")
VIZ_DIR = os.path.join(PROJECT, "visualizations")

doc = Document()

# ── Page setup: A4, narrow margins ──
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ── Default font ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.space_before = Pt(2)
style.paragraph_format.line_spacing = 1.15

# ── Page numbering from 22 ──
section = doc.sections[0]
sectPr = section._sectPr
pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="22"/>')
sectPr.append(pgNumType)

footer = section.footer
footer.is_linked_to_previous = False
fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
run2 = fp.add_run()
run2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
run3 = fp.add_run()
run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt({1: 14, 2: 13, 3: 12}[level])
    return h

def add_body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_code(code_text, label=""):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.italic = True
        r.font.color.rgb = RGBColor(80, 80, 80)
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)

def add_figure(img_path, caption, width=5.5):
    if not os.path.exists(img_path):
        add_body(f"[Image not found: {os.path.basename(img_path)}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = True


# ═══════════════════════════════════════════════════
# PAGE 1-2: Code Excerpts (concise, key methods only)
# ═══════════════════════════════════════════════════

add_heading("Appendix: Key Source Code & Working Prototype", level=1)
add_body(
    "This appendix presents the core implementation excerpts and working prototype "
    "screenshots of the MSME Viability Assessment System.",
    size=11
)

add_heading("A. Core Source Code Excerpts", level=2)

# A.1 SHAP
add_heading("A.1 SHAP Explainability Engine", level=3)
add_body(
    "Real-time feature contribution analysis using SHAP TreeExplainer, identifying "
    "the top 3 positive and negative factors for each prediction.",
    size=11
)
add_code("""
def explain(self, app_dict: dict):
    X = self._to_array(app_dict)
    pred = int(self.primary_model.predict(X)[0])
    sv_raw = self.shap_explainer.shap_values(X)

    if isinstance(sv_raw, list):
        sv = sv_raw[pred][0]
    elif isinstance(sv_raw, np.ndarray) and sv_raw.ndim == 3:
        sv = sv_raw[0, :, pred]
    else:
        sv = sv_raw[0]

    contributions = {self.feature_names[i]: float(sv[i])
                     for i in range(len(self.feature_names))}
    sorted_feats = sorted(contributions.items(),
                          key=lambda x: abs(x[1]), reverse=True)
    top_pos = [f for f, v in sorted_feats if v > 0][:3]
    top_neg = [f for f, v in sorted_feats if v < 0][:3]
    return {"predicted_class": pred,
            "feature_contributions": contributions,
            "top_positive_features": top_pos,
            "top_negative_features": top_neg}
""", label="File: api/engine.py — SHAP Explainability")

# A.2 Counterfactual
add_heading("A.2 Counterfactual Recommendation Engine", level=3)
add_body(
    "DiCE-inspired deterministic grid-search finding minimum feature changes "
    "to upgrade a business to a higher viability class.",
    size=11
)
add_code("""
FEATURES_TO_VARY = ["Term", "DisbursementGross", "SBA_Appv",
                     "GrAppv", "CreateJob", "RetainedJob"]

def recommend(self, app_dict, target_class=None):
    current_class = int(self.primary_model.predict(
        self._to_array(app_dict))[0])
    if target_class is None:
        target_class = min(current_class + 1, 4)

    perturbations = {
        "Term": [12, 24, 36, 48, 60, 84, 120, 180, 240],
        "CreateJob": list(range(1, 25)),
        "DisbursementGross": [0.9, 1.1, 1.2, 1.5, 2.0],
        "SBA_Appv": [0.9, 1.1, 1.2, 1.5, 2.0],
    }
    best_changes, best_dist = [], float("inf")

    # Single-feature perturbations first
    for feat in FEATURES_TO_VARY:
        for delta in perturbations[feat]:
            candidate = app_dict.copy()
            if feat in ["Term","CreateJob","RetainedJob"]:
                candidate[feat] += delta
            else:
                candidate[feat] *= delta
            pred = int(self.primary_model.predict(
                self._to_array(candidate))[0])
            if pred >= target_class:
                dist = abs(candidate[feat] - app_dict[feat])
                if dist < best_dist:
                    best_dist = dist
                    best_changes = [(feat, app_dict[feat],
                                     candidate[feat])]

    return {"current_label": LABEL_NAMES[current_class],
            "target_label": LABEL_NAMES[target_class],
            "feasible": len(best_changes) > 0,
            "changes": best_changes}
""", label="File: api/engine.py — Counterfactual Recommendations")

# A.3 Red Flags
add_heading("A.3 Red-Flag Detection Engine", level=3)
add_body(
    "Rule-based risk detection derived from 899,164 historical loan records.",
    size=11
)
add_code("""
def detect_red_flags(features, similar_data=None):
    flags = []
    # New business + large loan (28% default vs 15%)
    if (features.get("NewExist") == 2 and
        features.get("DisbursementGross", 0) > 100000):
        flags.append({"severity": "high",
            "flag": "High-value loan for new business",
            "explanation": "28% default vs 15% established.",
            "suggestion": "Start smaller, scale after 2-3 yrs."})

    # Heavy EMI burden
    term = features.get("Term", 84)
    amount = features.get("DisbursementGross", 0)
    if term > 0 and amount/term > 5000 and term < 60:
        flags.append({"severity": "high",
            "flag": "Heavy monthly EMI burden",
            "suggestion": f"Extend to {max(84,term*2)} months."})

    # Low guarantee coverage (<50%)
    sba = features.get("SBA_Appv", 0)
    gr = features.get("GrAppv", 1)
    if gr > 0 and sba/gr < 0.5:
        flags.append({"severity": "medium",
            "flag": f"Low guarantee ({sba/gr*100:.0f}%)",
            "suggestion": "Apply for CGTMSE (up to 85%)."})
    return sorted(flags, key=lambda x:
        {"high":0,"medium":1,"low":2}[x["severity"]])
""", label="File: api/optimizer.py — Red-Flag Detection")

# A.4 Loan Optimizer
add_heading("A.4 Loan Structure Optimizer", level=3)
add_body(
    "Grid-search across 12 term lengths and 10 amount multipliers to find "
    "the safest loan configuration via the ML model.",
    size=11
)
add_code("""
class LoanOptimizer:
    def __init__(self, prediction_engine):
        self.engine = prediction_engine

    def find_optimal_term(self, features):
        terms = [12,24,36,48,60,72,84,120,180,240,300,360]
        results = []
        for term in terms:
            test = deepcopy(features)
            test["Term"] = term
            pred = self.engine.predict(test)
            results.append({"term": term,
                "class": pred["predicted_label"],
                "confidence": pred["confidence"]})
        best = max(results, key=lambda x:
            (x["class"], x["confidence"]))
        return {"recommended_term": best["term"],
                "best_class": best["class"]}

    def find_max_safe_amount(self, features, target=2):
        original = features.get("DisbursementGross", 100000)
        for mult in [0.1,0.25,0.5,1.0,1.5,2.0,3.0,5.0]:
            test = deepcopy(features)
            test["DisbursementGross"] = int(original * mult)
            pred = self.engine.predict(test)
            if pred["predicted_class"] >= target:
                max_safe = int(original * mult)
        return {"max_safe_amount": max_safe,
                "can_take_more": max_safe > original}
""", label="File: api/optimizer.py — Loan Structure Optimizer")


# ═══════════════════════════════════════════════════
# PAGE 3: Visualizations (5 figures)
# ═══════════════════════════════════════════════════

add_heading("B. Model Training Results & Visualizations", level=2)

add_figure(
    os.path.join(VIZ_DIR, "model_comparison_table.png"),
    "Figure B.1: Comparative performance of all classifiers on the "
    "5-class viability spectrum (test set n=178,385).",
    width=5.8
)

add_figure(
    os.path.join(VIZ_DIR, "xgb_confusion_matrix.png"),
    "Figure B.2: XGBoost 5-class confusion matrix — 92.78% accuracy.",
    width=3.8
)

add_figure(
    os.path.join(VIZ_DIR, "multiclass_distribution.png"),
    "Figure B.3: 5-class health label distribution across 899,164 records.",
    width=4.2
)

add_figure(
    os.path.join(VIZ_DIR, "mc_shap_global_bar.png"),
    "Figure B.4: SHAP global feature importance — Term, DisbursementGross, "
    "and SBA_Appv dominate predictions.",
    width=4.5
)

add_figure(
    os.path.join(VIZ_DIR, "shap_summary_plot.png"),
    "Figure B.5: SHAP beeswarm plot — per-instance feature contributions "
    "colored by feature value.",
    width=4.5
)


# ═══════════════════════════════════════════════════
# PAGE 4+: Working Prototype Screenshots (7 figures)
# ═══════════════════════════════════════════════════

add_heading("C. Working Prototype Screenshots", level=2)
add_body(
    "The following screenshots demonstrate the complete assessment flow of the "
    "deployed MSME Loan Readiness Coach, from parameter input through viability "
    "grading, prescriptive recommendations, and government scheme matching.",
    size=11
)

# C.1 Input form + grade
add_figure(
    os.path.join(SS_DIR, "ss1_input_and_grade.png"),
    "Figure C.1: Expert Mode — loan parameters entered, assessment submitted. "
    "Result shows Grade C (Stable) at 99.6% confidence with business radar chart.",
    width=5.5
)

# C.2 Grade + radar detail
add_figure(
    os.path.join(SS_DIR, "ss2_grade_radar.png"),
    "Figure C.2: Viability grade card showing 'C — Stable' with radar chart "
    "visualizing the business profile across Maturity, Employment, Loan Term, "
    "Documentation, Location, and Guarantee dimensions.",
    width=5.5
)

# C.3 Recommendations
add_figure(
    os.path.join(SS_DIR, "ss3_recommendations.png"),
    "Figure C.3: Prescriptive recommendations — the counterfactual engine suggests "
    "extending term from 84 to 240 months, switching to full documentation, and "
    "increasing the guarantee amount to upgrade from Stable → Growing (99% confidence).",
    width=5.5
)

# C.4 Similar businesses + Govt schemes
add_figure(
    os.path.join(SS_DIR, "ss4_similar_biz_schemes.png"),
    "Figure C.4: Peer comparison (50 similar businesses, 100% success rate vs 82% baseline) "
    "and auto-matched government schemes — MUDRA Yojana, NSIC, and CGTMSE.",
    width=5.5
)

# C.5 SHAP analysis
add_figure(
    os.path.join(SS_DIR, "ss5_shap_analysis.png"),
    "Figure C.5: Real-time SHAP feature contribution chart showing Term and SBA_Appv "
    "as the strongest positive contributors to the Stable classification.",
    width=5.5
)

# C.6 AI Coach
add_figure(
    os.path.join(SS_DIR, "05_loan_coach_chat.png"),
    "Figure C.6: AI Loan Coach — conversational assessment interface with Hindi/English "
    "support, voice input capability, and natural language business description.",
    width=5.0
)

# C.7 Analytics
add_figure(
    os.path.join(SS_DIR, "06_analytics_dashboard.png"),
    "Figure C.7: Analytics Dashboard — 85 predictions logged in SQLite audit trail, "
    "89.5% average confidence, 4 risk classes observed, with distribution histogram.",
    width=5.0
)


# ── Save ──
out_path = os.path.join(PROJECT, "Appendix_Source_Code_and_Prototype.docx")
doc.save(out_path)
print(f"✅ Created: {out_path}")
print(f"   Size: {os.path.getsize(out_path) / 1024:.0f} KB")
